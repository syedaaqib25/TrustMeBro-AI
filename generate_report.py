"""
Generate Seminar Report Word Document
CSE435 - Comprehensive Seminar | CA2
Lovely Professional University Punjab
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Config ──────────────────────────────────────────────────────────
STUDENT_NAME = "Syeda Aqib"
REG_NO = "XXXXXXXXXX"         # ← Replace with your actual Reg No
FACULTY_NAME = "Dr. XXXX XXXX"  # ← Replace with your supervisor's name
FACULTY_DESIGNATION = "Assistant Professor"
ACADEMIC_SESSION = "Jan–Apr 2026"
MONTH_YEAR = "April, 2026"
GITHUB_LINK = "https://github.com/syedaaqib25/TrustMeBro-AI"
LINKEDIN_LINK = "https://linkedin.com/in/your-profile"  # ← Replace

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Seminar_Report_TrustMeBro_AI.docx")


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    """Add page numbers at bottom center."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char1)
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)
        run3 = p.add_run()
        fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char2)
        for r in [run, run2, run3]:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def heading(doc, text, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_heading(text, level=level)
    p.alignment = alignment
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._r.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
        rPr.insert(0, rFonts)
    return p


def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return run


def normal_run(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return run


def para(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    return p


def bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        bold_run(p, bold_prefix)
        normal_run(p, text)
    else:
        run = p.runs[0] if p.runs else p.add_run(text)
        if not p.runs:
            pass
        else:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
        if not p.text:
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def empty_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_normal_style(doc)

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    empty_lines(doc, 3)
    para(doc, "Seminar Report", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "On", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TrustMeBro — Multi-Model Fake News Credibility Analyzer")
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    empty_lines(doc, 2)
    para(doc, "Submitted by", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{STUDENT_NAME}          Registration No. {REG_NO}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 1)
    para(doc, "Bachelor of Technology", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "IN", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Computer Science and Engineering", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 1)
    para(doc, "Under the Supervision of", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, FACULTY_NAME, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, FACULTY_DESIGNATION, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LOVELY PROFESSIONAL UNIVERSITY PUNJAB")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    para(doc, f"({MONTH_YEAR})", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # DECLARATION PAGE
    # ══════════════════════════════════════════════════════════════
    empty_lines(doc, 2)
    heading(doc, "DECLARATION", level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 1)

    para(doc, (
        f'I hereby declare that the seminar report titled "TrustMeBro — Multi-Model Fake News '
        f'Credibility Analyzer" submitted in partial fulfillment of the requirements for the award '
        f'of the degree of Bachelor of Technology in Computer Science and Engineering is a record of '
        f'my own work carried out during the academic session {ACADEMIC_SESSION}.'
    ))

    para(doc, (
        "I further declare that this report has not been submitted, either in part or in full, to any "
        "other institution or university for the award of any degree or diploma."
    ))

    para(doc, (
        "I confirm that the content of this report is original and prepared by me. Any references used "
        "have been duly acknowledged. I also declare that the use of Artificial Intelligence (AI) tools, "
        "if any, has been minimal and the AI-generated content in this report is less than 10%, ensuring "
        "that the majority of the work reflects my own understanding and effort."
    ))

    para(doc, "I take full responsibility for the authenticity and originality of the work presented in this report.")

    empty_lines(doc, 2)
    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "Name of the Student: ")
    normal_run(p, STUDENT_NAME)

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "Registration Number: ")
    normal_run(p, REG_NO)

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "Course: ")
    normal_run(p, "B.Tech CSE")

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "Signature of the Student: ")
    normal_run(p, "____________________")

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "Date: ")
    normal_run(p, "____________________")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ══════════════════════════════════════════════════════════════
    heading(doc, "TABLE OF CONTENTS", level=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    empty_lines(doc, 1)

    toc_items = [
        ("Chapter 1", "Introduction", "1"),
        ("", "1.1  Title of the Seminar Topic", "1"),
        ("", "1.2  Background and Importance", "1"),
        ("", "1.3  Objectives of the Seminar", "2"),
        ("", "1.4  Brief Overview of Methodology", "3"),
        ("Chapter 2", "Literature Review", "4"),
        ("", "2.1  Existing Approaches to Fake News Detection", "4"),
        ("", "2.2  Machine Learning in News Classification", "5"),
        ("", "2.3  Role of Large Language Models", "6"),
        ("", "2.4  Research Gaps", "6"),
        ("Chapter 3", "Conceptual Study / Seminar Work", "7"),
        ("", "3.1  System Architecture", "7"),
        ("", "3.2  Data Pipeline", "8"),
        ("", "3.3  Feature Engineering", "9"),
        ("", "3.4  Model Descriptions", "10"),
        ("", "3.5  AI Integration (Gemini & ChatGPT)", "11"),
        ("", "3.6  Tools and Technologies", "12"),
        ("Chapter 4", "Results and Discussion", "13"),
        ("", "4.1  Model Performance Metrics", "13"),
        ("", "4.2  Comparative Analysis", "14"),
        ("", "4.3  Interpretation and Insights", "15"),
        ("", "4.4  Advantages and Limitations", "16"),
        ("Chapter 5", "Conclusion and Future Scope", "17"),
        ("", "5.1  Summary of Work", "17"),
        ("", "5.2  Major Learning Outcomes", "17"),
        ("", "5.3  Conclusions", "18"),
        ("", "5.4  Future Scope", "18"),
        ("", "Professional Profile & Repository Details", "19"),
        ("", "References", "20"),
    ]
    for chap, title, pg in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.tab_stops.add_tab_stop(Inches(5.75), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        if chap:
            bold_run(p, f"{chap}: ")
        normal_run(p, f"{title}\t{pg}")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 1: INTRODUCTION
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Chapter-1: Introduction", level=1)

    heading(doc, "1.1  Title of the Seminar Topic", level=2)
    para(doc, (
        "TrustMeBro — Multi-Model Fake News Credibility Analyzer: A next-generation news "
        "credibility analysis platform that combines classical machine learning models with "
        "modern large language models to deliver real-time, multi-faceted verdicts on news articles."
    ))

    heading(doc, "1.2  Background and Importance of the Topic", level=2)
    para(doc, (
        "The proliferation of misinformation and fake news across digital media platforms has become "
        "one of the most pressing challenges of the modern information age. With the rapid growth of "
        "social media, online news outlets, and user-generated content, distinguishing credible news "
        "from fabricated or misleading articles has become increasingly difficult for the general public. "
        "Studies have shown that fake news spreads six times faster than true news on platforms like "
        "Twitter, and can have severe consequences on public opinion, democratic processes, financial "
        "markets, and public health — as evidenced during the COVID-19 pandemic."
    ))
    para(doc, (
        "Traditional fact-checking methods, while valuable, are inherently slow and cannot keep pace "
        "with the volume of content being generated daily. This creates a critical need for automated "
        "systems that can analyze news articles in real-time and provide credibility assessments. Machine "
        "learning and natural language processing (NLP) techniques have shown significant promise in "
        "addressing this challenge, with models achieving over 99% accuracy on benchmark datasets."
    ))
    para(doc, (
        "Furthermore, the emergence of Large Language Models (LLMs) such as Google Gemini and OpenAI's "
        "GPT-4 has opened new avenues for zero-shot reasoning about content credibility, providing "
        "interpretable explanations alongside their assessments. The combination of high-accuracy ML "
        "classifiers with the reasoning capabilities of LLMs represents a powerful hybrid approach that "
        "this project explores."
    ))

    heading(doc, "1.3  Objectives of the Seminar", level=2)
    objectives = [
        "To study and implement multiple machine learning algorithms for binary text classification of news articles as credible or fake.",
        "To design and develop a complete data pipeline for preprocessing, feature extraction (TF-IDF), and model training on a labeled news dataset of 23,941 samples.",
        "To integrate two Large Language Models — Google Gemini 2.0 Flash and OpenAI GPT-4o-mini — for zero-shot credibility reasoning alongside classical ML models.",
        "To build a production-ready REST API using FastAPI that runs all 6 analysis signals (4 ML + 2 AI) in parallel with sub-second response times.",
        "To develop an interactive, responsive web dashboard featuring article analysis, live news feeds, analysis history, user feedback, and an admin panel.",
        "To evaluate and compare model performances using standard metrics — Accuracy, Precision, Recall, F1-Score, and ROC-AUC.",
        "To implement a weighted ensemble scoring formula that aggregates diverse model outputs into a unified 0–100 credibility score.",
    ]
    for obj in objectives:
        bullet(doc, obj)

    heading(doc, "1.4  Brief Overview of the Approach / Methodology", level=2)
    para(doc, (
        "The TrustMeBro system follows a systematic methodology encompassing data collection, "
        "preprocessing, feature engineering, model training, API development, and frontend design. "
        "The approach can be summarized in the following phases:"
    ))
    phases = [
        ("Data Collection & Preprocessing: ", "Two labeled CSV datasets (True.csv and Fake.csv) are combined, cleaned using regex-based text normalization, NLTK tokenization, stopword removal, and WordNet lemmatization."),
        ("Feature Engineering: ", "TF-IDF vectorization with up to 50,000 features and unigram-bigram support transforms text into numerical representations suitable for ML classifiers."),
        ("Model Training: ", "Four classical ML models — Logistic Regression, Support Vector Machine (SVM), Naive Bayes, and LightGBM — are trained on the TF-IDF features with class-weight balancing and cross-validation."),
        ("AI Integration: ", "Google Gemini 2.0 Flash and OpenAI GPT-4o-mini are queried asynchronously to provide zero-shot credibility scores and reasoning."),
        ("Parallel Execution: ", "FastAPI with ThreadPoolExecutor runs all 6 models concurrently, aggregating results into a weighted overall credibility score."),
        ("Frontend Dashboard: ", "A responsive single-page application built with vanilla JavaScript and custom CSS provides the user interface."),
    ]
    for bold_text, desc in phases:
        bullet(doc, desc, bold_prefix=bold_text)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 2: LITERATURE REVIEW
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Chapter-2: Literature Review", level=1)

    heading(doc, "2.1  Existing Approaches to Fake News Detection", level=2)
    para(doc, (
        "Fake news detection has garnered significant research attention over the past decade. "
        "Early approaches relied on linguistic feature analysis, examining stylistic markers such as "
        "sensationalist language, excessive use of capital letters, and emotional tone. Researchers "
        "like Pérez-Rosas et al. (2018) demonstrated that linguistic cues alone could achieve moderate "
        "accuracy in distinguishing fake from real news, but these methods struggled with sophisticated "
        "misinformation that mimics credible writing styles."
    ))
    para(doc, (
        "Network-based approaches analyzed the propagation patterns of news articles on social media. "
        "Vosoughi et al. (2018) published a seminal study in Science showing that false news spreads "
        "significantly faster and reaches more people than true news. However, these methods require "
        "social network data that is not always available at the time of article publication."
    ))
    para(doc, (
        "Knowledge-graph-based approaches attempted to verify factual claims against curated databases. "
        "While highly precise, they are limited by the coverage of existing knowledge bases and cannot "
        "handle novel or emerging topics."
    ))

    heading(doc, "2.2  Machine Learning in News Classification", level=2)
    para(doc, (
        "Machine learning approaches to text classification have evolved significantly. Traditional "
        "methods such as Naive Bayes and Logistic Regression, when combined with TF-IDF features, "
        "have proven surprisingly effective for fake news detection. Ahmed et al. (2017) reported "
        "accuracies exceeding 92% using Linear SVM with TF-IDF on news datasets."
    ))
    para(doc, (
        "Ensemble methods and gradient boosting frameworks like XGBoost and LightGBM have further "
        "improved performance by combining multiple weak learners. Raza and Ding (2022) showed that "
        "ensemble approaches consistently outperform individual classifiers for fake news detection, "
        "particularly when dealing with imbalanced datasets."
    ))
    para(doc, (
        "Deep learning models, including Convolutional Neural Networks (CNNs) and Long Short-Term "
        "Memory (LSTM) networks, have been applied to capture sequential patterns in text. While "
        "they can model complex relationships, they often require significantly more data and "
        "computational resources than classical methods, sometimes without proportional accuracy gains."
    ))

    heading(doc, "2.3  Role of Large Language Models", level=2)
    para(doc, (
        "The advent of Large Language Models has introduced a paradigm shift in NLP tasks. Models "
        "like GPT-4, Gemini, and Claude can perform zero-shot classification — analyzing text without "
        "being explicitly trained on labeled fake news data. Their ability to reason about content, "
        "identify logical inconsistencies, and assess source credibility makes them valuable complement "
        "to traditional ML classifiers."
    ))
    para(doc, (
        "Recent studies by Huang et al. (2023) explored LLMs as fact-checkers and found that while "
        "they may not match fine-tuned models on benchmark accuracy, they provide interpretable "
        "explanations that are highly valuable for end-users. The combination of LLM reasoning with "
        "ML precision is an emerging area of active research."
    ))

    heading(doc, "2.4  Research Gaps Identified", level=2)
    para(doc, (
        "The literature review reveals several gaps that the TrustMeBro system aims to address:"
    ))
    gaps = [
        "Most existing systems rely on a single model or model type, lacking multi-signal cross-validation.",
        "Few systems combine classical ML with LLMs in a parallel ensemble architecture.",
        "Real-time analysis with sub-second latency is rarely addressed in academic prototypes.",
        "User interaction features such as feedback loops for model improvement are largely absent.",
        "There is limited focus on providing interpretable reasoning alongside credibility scores.",
    ]
    for g in gaps:
        bullet(doc, g)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 3: CONCEPTUAL STUDY / SEMINAR WORK
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Chapter-3: Conceptual Study / Seminar Work", level=1)

    heading(doc, "3.1  System Architecture", level=2)
    para(doc, (
        "The TrustMeBro system employs a modular, layered architecture that separates concerns across "
        "data processing, machine learning, AI integration, API serving, and frontend presentation. "
        "The architecture consists of the following key components:"
    ))
    bullet(doc, "A FastAPI-based REST backend that handles HTTP requests, routing, and CORS middleware.")
    bullet(doc, "A ThreadPoolExecutor that runs 4 ML models concurrently on incoming article text.")
    bullet(doc, "Async integration with Google Gemini 2.0 Flash and OpenAI GPT-4o-mini for zero-shot analysis.")
    bullet(doc, "A data pipeline for CSV ingestion, NLTK-based preprocessing, and TF-IDF vectorization.")
    bullet(doc, "A single-page frontend application served as static files by FastAPI.")

    para(doc, (
        "Figure 3.1 illustrates the high-level system architecture showing the request flow from the "
        "user through the API layer to the parallel model execution engine."
    ))

    # Architecture flow description
    para(doc, "Request Flow:", bold=True)
    para(doc, (
        "When a user submits an article for analysis, the frontend sends a POST request to the /analyze "
        "endpoint. The API layer simultaneously dispatches the text to all 4 ML models via a thread pool "
        "and to both AI models via async coroutines. Upon receiving all 6 responses, the system computes "
        "a weighted average credibility score and returns a comprehensive JSON response containing "
        "individual model scores, AI reasoning, and the overall verdict."
    ))

    heading(doc, "3.2  Data Pipeline", level=2)
    para(doc, (
        "The data pipeline is responsible for transforming raw CSV data into clean, vectorized features "
        "ready for model training. The pipeline is implemented in the data_pipeline.py module and "
        "consists of the following stages:"
    ))

    para(doc, "Table 3.1: Data Pipeline Stages", bold=True)
    add_table(doc,
        ["Stage", "Operation", "Details"],
        [
            ["1. Load", "CSV Ingestion", "Read True.csv and Fake.csv, assign labels (1/0)"],
            ["2. Combine", "Text Merging", "Concatenate title + text into single column"],
            ["3. Clean", "Normalization", "Lowercase, remove URLs, HTML, special characters"],
            ["4. Tokenize", "Word Splitting", "NLTK word_tokenize, stopword removal"],
            ["5. Lemmatize", "Root Form", "WordNet lemmatizer for morphological normalization"],
            ["6. Vectorize", "TF-IDF", "Up to 50,000 features, unigrams + bigrams, sublinear TF"],
            ["7. Split", "Train/Test", "80/20 stratified split with random_state=42"],
        ]
    )

    empty_lines(doc, 1)
    para(doc, (
        "The preprocessing pipeline handles edge cases including empty strings, non-string inputs, "
        "and multi-space normalization. NLTK resources are downloaded lazily with timeout protection "
        "and force-loaded to prevent threading race conditions during parallel processing."
    ))

    heading(doc, "3.3  Feature Engineering", level=2)
    para(doc, (
        "The primary feature extraction method used is Term Frequency-Inverse Document Frequency (TF-IDF) "
        "vectorization, implemented using scikit-learn's TfidfVectorizer. The configuration parameters are:"
    ))
    bullet(doc, "50,000 features (reduced from vocabulary for efficiency)", bold_prefix="max_features: ")
    bullet(doc, "(1, 2) — captures both individual words and two-word phrases", bold_prefix="ngram_range: ")
    bullet(doc, "True — applies logarithmic scaling to term frequencies", bold_prefix="sublinear_tf: ")
    para(doc, (
        "The resulting feature matrix has 32,774 features after fitting on the training set of 23,941 "
        "samples. The fitted vectorizer is serialized as tfidf_vectorizer.pkl for reuse during prediction."
    ))
    para(doc, (
        "Additionally, the system supports Word2Vec embeddings (200-dimensional, trained via Gensim) "
        "and BERT [CLS] token embeddings (768-dimensional) for deep learning and transformer-based "
        "models respectively."
    ))

    heading(doc, "3.4  Model Descriptions", level=2)

    para(doc, "3.4.1  Naive Bayes (MultinomialNB)", bold=True)
    para(doc, (
        "A probabilistic classifier based on Bayes' theorem with strong independence assumptions between "
        "features. Configured with a smoothing parameter alpha=0.1, it serves as a fast baseline model. "
        "Despite its simplicity, it achieves 96.52% accuracy on the test set."
    ))

    para(doc, "3.4.2  Logistic Regression", bold=True)
    para(doc, (
        "A linear classifier optimized with L2 regularization (C=1.0), balanced class weights, and a "
        "maximum of 1,000 iterations. It models the probability of an article being true or fake as a "
        "logistic function of TF-IDF features, achieving 99.15% accuracy."
    ))

    para(doc, "3.4.3  Support Vector Machine (SVM)", bold=True)
    para(doc, (
        "A LinearSVC model wrapped in CalibratedClassifierCV (3-fold) to produce probability estimates. "
        "SVMs find the optimal hyperplane that maximizes the margin between classes in the high-dimensional "
        "TF-IDF feature space, achieving 99.62% accuracy."
    ))

    para(doc, "3.4.4  LightGBM", bold=True)
    para(doc, (
        "A gradient boosting framework using tree-based learning algorithms. Configured with 100 estimators, "
        "balanced class weights, and wrapped in CalibratedClassifierCV for probability calibration. It "
        "achieves the highest accuracy of 99.69% among all ML models, leveraging feature importance based "
        "on information gain."
    ))

    heading(doc, "3.5  AI Integration (Gemini & ChatGPT)", level=2)
    para(doc, (
        "The system integrates two state-of-the-art Large Language Models for zero-shot credibility analysis:"
    ))

    para(doc, "Google Gemini 2.0 Flash:", bold=True)
    para(doc, (
        "Called via the google-genai Python client, Gemini receives the first 2,000 characters of the article "
        "with a structured prompt requesting a 0–100 credibility score and one-sentence reasoning. The response "
        "is parsed from JSON format. Rate limiting is enforced at 14 requests per minute and 1,400 per day "
        "to stay within free-tier quotas."
    ))

    para(doc, "OpenAI GPT-4o-mini:", bold=True)
    para(doc, (
        "Called via the OpenAI Python SDK, GPT-4o-mini receives the same structured prompt with temperature=0.3 "
        "for consistent outputs and max_tokens=150. Both AI models include automatic retry with backoff on "
        "429/quota errors and friendly user-facing error messages."
    ))

    para(doc, "Weighted Ensemble Formula:", bold=True)
    para(doc, (
        "When both AI models respond successfully, the overall score is calculated as: "
        "Overall = (ML_average × 0.50) + (Gemini_score × 0.25) + (ChatGPT_score × 0.25). "
        "If only one AI model is available, the split becomes 60% ML / 40% AI. If neither is available, "
        "the score defaults to 100% ML average."
    ))

    heading(doc, "3.6  Tools, Platforms, and Technologies", level=2)
    para(doc, "Table 3.2: Technology Stack", bold=True)
    add_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Backend", "FastAPI, Uvicorn, asyncio", "REST API, async request handling"],
            ["ML Models", "scikit-learn, LightGBM", "Text classification (TF-IDF → classifiers)"],
            ["NLP", "NLTK, Regex", "Text preprocessing, tokenization, stopwords"],
            ["AI", "Google Gemini 2.0 Flash", "Zero-shot credibility reasoning"],
            ["AI", "OpenAI GPT-4o-mini", "Zero-shot credibility reasoning"],
            ["Data", "Pandas, NumPy", "Data manipulation and pipeline"],
            ["Scraping", "BeautifulSoup4, Requests", "URL article extraction"],
            ["News", "NewsAPI.org", "Live headline feed"],
            ["Frontend", "Vanilla JS (ES6+)", "Interactive SPA"],
            ["Styling", "Custom CSS", "Glassmorphism, animations, responsive grid"],
            ["Testing", "pytest, TestClient", "25 integration tests"],
        ]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 4: RESULTS AND DISCUSSION
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Chapter-4: Results and Discussion", level=1)

    heading(doc, "4.1  Model Performance Metrics", level=2)
    para(doc, (
        "All four classical ML models were trained on 19,152 samples (80% split) and evaluated on "
        "4,789 samples (20% split) from the cleaned dataset of 23,941 news articles. The following "
        "table summarizes the key performance metrics:"
    ))

    para(doc, "Table 4.1: Model Performance Comparison", bold=True)
    add_table(doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"],
        [
            ["LightGBM", "99.69%", "99.53%", "99.81%", "99.67%", "99.98%"],
            ["SVM", "99.62%", "99.53%", "99.67%", "99.60%", "99.99%"],
            ["Logistic Regression", "99.15%", "98.86%", "99.37%", "99.12%", "99.95%"],
            ["Naive Bayes", "96.52%", "96.44%", "96.26%", "96.35%", "99.33%"],
        ]
    )

    empty_lines(doc, 1)
    para(doc, (
        "All models demonstrate exceptional performance, with LightGBM achieving the highest accuracy "
        "of 99.69% and F1-score of 99.67%. The SVM model closely follows with 99.62% accuracy and "
        "the highest ROC-AUC of 99.99%. Even the baseline Naive Bayes model achieves a respectable "
        "96.52% accuracy."
    ))

    heading(doc, "4.2  Comparative Analysis", level=2)
    para(doc, (
        "Key observations from the comparative analysis of model performances:"
    ))
    bullet(doc, "LightGBM has the highest recall (99.81%), making it the most sensitive model — it rarely misses fake news articles.", bold_prefix="Best Recall: ")
    bullet(doc, "SVM achieves the highest ROC-AUC (99.99%), indicating near-perfect discrimination capability across all classification thresholds.", bold_prefix="Best ROC-AUC: ")
    bullet(doc, "All four models share the same precision (99.53%) for the top three, suggesting robust false-positive control.", bold_prefix="Consistent Precision: ")
    bullet(doc, "While Naive Bayes underperforms relative to others, its 96.52% accuracy and sub-millisecond inference time make it valuable for resource-constrained environments.", bold_prefix="Naive Bayes Trade-off: ")

    para(doc, (
        "The AI models (Gemini 2.0 Flash and GPT-4o-mini) do not have traditional ML metrics since "
        "they perform zero-shot analysis. However, their scores contribute to the weighted ensemble "
        "and provide interpretable reasoning that enhances user trust in the system's verdicts."
    ))

    heading(doc, "4.3  Interpretation and Insights", level=2)
    para(doc, (
        "The exceptionally high accuracy across all models can be attributed to several factors:"
    ))
    bullet(doc, "TF-IDF with bigrams captures both individual word importance and phrase-level patterns that are strong indicators of credibility (e.g., sensationalist phrases common in fake news).")
    bullet(doc, "The dataset exhibits clear stylistic differences between true and fake news articles — fake articles tend to use more emotional language, lack specific details, and follow different structural patterns.")
    bullet(doc, "Class-weight balancing and calibrated classifiers prevent bias toward the majority class and ensure reliable probability estimates.")
    bullet(doc, "Sublinear TF scaling prevents long documents from dominating the feature space, leading to more balanced predictions.")

    para(doc, "Table 4.2: Credibility Score Interpretation", bold=True)
    add_table(doc,
        ["Score Range", "Label", "Interpretation"],
        [
            ["70–100", "Likely Credible", "Article exhibits characteristics of credible news"],
            ["40–69", "Uncertain", "Article has mixed signals; manual verification recommended"],
            ["0–39", "Likely Fake", "Article shows strong indicators of fabricated content"],
        ]
    )

    heading(doc, "4.4  Advantages and Limitations", level=2)
    para(doc, "Advantages:", bold=True)
    bullet(doc, "Multi-signal cross-validation using 6 independent analysis channels reduces false positives.")
    bullet(doc, "Parallel execution delivers results in under 2 seconds even with AI model calls.")
    bullet(doc, "Interpretable AI reasoning helps users understand why an article is flagged.")
    bullet(doc, "User feedback system enables continuous improvement through data collection.")
    bullet(doc, "Live headline integration allows one-click analysis of trending news.")

    para(doc, "Limitations:", bold=True)
    bullet(doc, "The ML models are trained on English-language news only and may not generalize to other languages.")
    bullet(doc, "AI model scores depend on API availability and are subject to rate limits and costs.")
    bullet(doc, "The system analyzes text content only — it does not verify claims against external fact-check databases.")
    bullet(doc, "Novel misinformation tactics that deviate from training data patterns may evade detection.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 5: CONCLUSION AND FUTURE SCOPE
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Chapter-5: Conclusion and Future Scope", level=1)

    heading(doc, "5.1  Summary of the Seminar Work", level=2)
    para(doc, (
        "This seminar presents TrustMeBro, a comprehensive fake news credibility analysis platform that "
        "combines the precision of classical machine learning with the reasoning capabilities of modern "
        "Large Language Models. The system was developed as a full-stack application featuring a FastAPI "
        "backend, 4 trained ML classifiers (Logistic Regression, SVM, Naive Bayes, LightGBM), 2 AI "
        "integrations (Gemini 2.0 Flash and GPT-4o-mini), and a responsive web dashboard."
    ))
    para(doc, (
        "The data pipeline processes 23,941 labeled news articles through comprehensive text cleaning, "
        "NLTK-based tokenization and lemmatization, and TF-IDF vectorization with 32,774 features. "
        "All models were trained with stratified splitting, class-weight balancing, and cross-validated "
        "calibration to ensure reliable probability estimates."
    ))

    heading(doc, "5.2  Major Learning Outcomes", level=2)
    bullet(doc, "Gained hands-on experience with end-to-end ML pipeline development — from data preprocessing to model deployment.")
    bullet(doc, "Learned to build asynchronous, concurrent Python backends using FastAPI, asyncio, and ThreadPoolExecutor.")
    bullet(doc, "Understood the practical integration of LLM APIs (Gemini, OpenAI) including rate limiting, error handling, and retry strategies.")
    bullet(doc, "Explored TF-IDF vectorization, its parameters, and its effectiveness for text classification tasks.")
    bullet(doc, "Developed skills in full-stack web development with vanilla JavaScript and modern CSS techniques.")
    bullet(doc, "Gained experience with software testing — 25 integration tests using pytest and FastAPI TestClient.")

    heading(doc, "5.3  Conclusions", level=2)
    para(doc, (
        "The key conclusions drawn from this seminar work are:"
    ))
    bullet(doc, "Classical ML models, when properly configured with TF-IDF features and class-weight balancing, can achieve over 99% accuracy for fake news detection on well-curated datasets.")
    bullet(doc, "A multi-model ensemble approach provides more robust and trustworthy predictions than any single model alone.")
    bullet(doc, "LLMs serve as excellent complementary signals — while their individual accuracy may vary, their reasoning capabilities add significant interpretive value.")
    bullet(doc, "The weighted ensemble formula effectively combines diverse model outputs into a meaningful credibility score that is actionable for end-users.")
    bullet(doc, "Parallel execution architecture is crucial for real-time analysis — running 6 models concurrently keeps response times under 2 seconds.")

    heading(doc, "5.4  Future Scope", level=2)
    para(doc, (
        "Several promising directions for future enhancement of this work include:"
    ))
    bullet(doc, "Fine-tuning transformer models (BERT/DistilBERT) on the dataset for improved accuracy and transfer learning capabilities.", bold_prefix="Transformer Fine-tuning: ")
    bullet(doc, "Integrating a source credibility database to cross-reference article sources with known reliability ratings.", bold_prefix="Source Verification: ")
    bullet(doc, "Extending the system to support news articles in Hindi, Spanish, French, and other languages.", bold_prefix="Multi-language Support: ")
    bullet(doc, "Developing a Chrome/Firefox extension for one-click credibility checks from any news website.", bold_prefix="Browser Extension: ")
    bullet(doc, "Implementing user accounts with personalized analysis history and saved preferences.", bold_prefix="User Accounts: ")
    bullet(doc, "Building an automated retraining pipeline that uses collected user feedback to improve model accuracy over time.", bold_prefix="Retraining Pipeline: ")
    bullet(doc, "Connecting to fact-checking APIs (e.g., ClaimBuster, Google Fact Check Tools) for claim-level verification.", bold_prefix="Fact-Check Integration: ")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # PROFESSIONAL PROFILE & REPOSITORY
    # ══════════════════════════════════════════════════════════════
    heading(doc, "Professional Profile & Repository Details", level=1)
    empty_lines(doc, 1)

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "GitHub Project Repository: ")
    normal_run(p, GITHUB_LINK)

    p = para(doc, "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    bold_run(p, "LinkedIn Profile: ")
    normal_run(p, LINKEDIN_LINK)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════
    heading(doc, "References", level=1)
    empty_lines(doc, 1)

    refs = [
        "Ahmed, H., Traore, I., & Saad, S. (2017). Detection of Online Fake News Using N-Gram Analysis and Machine Learning Techniques. Lecture Notes in Computer Science, 10618, 127–138.",
        "Vosoughi, S., Roy, D., & Aral, S. (2018). The spread of true and false news online. Science, 359(6380), 1146–1151.",
        "Pérez-Rosas, V., Kleinberg, B., Lefevre, A., & Mihalcea, R. (2018). Automatic Detection of Fake News. Proceedings of the 27th International Conference on Computational Linguistics, 3391–3401.",
        "Raza, S., & Ding, C. (2022). Fake news detection based on news content and social contexts: a transformer-based approach. International Journal of Data Science and Analytics, 13, 335–362.",
        "Huang, Y., Sun, L., & Qiu, X. (2023). Can Large Language Models Be Used for Fact-Checking? arXiv preprint arXiv:2308.04945.",
        "Shu, K., Sliva, A., Wang, S., Tang, J., & Liu, H. (2017). Fake News Detection on Social Media: A Data Mining Perspective. ACM SIGKDD Explorations Newsletter, 19(1), 22–36.",
        "Zhou, X., & Zafarani, R. (2020). A Survey of Fake News: Fundamental Theories, Detection Methods, and Opportunities. ACM Computing Surveys, 53(5), 1–40.",
        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "Ke, G., et al. (2017). LightGBM: A Highly Efficient Gradient Boosting Decision Tree. Advances in Neural Information Processing Systems, 30.",
        "Bird, S., Klein, E., & Loper, E. (2009). Natural Language Processing with Python. O'Reilly Media.",
        "FastAPI Documentation. https://fastapi.tiangolo.com/",
        "Google Gemini API Documentation. https://ai.google.dev/",
        "OpenAI API Documentation. https://platform.openai.com/docs/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(f"[{i}]  {ref}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    # ── Page Numbers ──
    add_page_number(doc)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Report generated successfully!")
    print(f"📄 File: {OUTPUT_PATH}")
    print(f"📦 Size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
    print(f"\n⚠️  Remember to update placeholder values:")
    print(f"    - Registration Number (currently: {REG_NO})")
    print(f"    - Faculty Name (currently: {FACULTY_NAME})")
    print(f"    - LinkedIn Profile URL")
    print(f"    - Update Table of Contents page numbers in Word")


if __name__ == "__main__":
    build()
